# -*- coding: utf-8 -*-

from odoo import models, fields, _
from datetime import datetime
from odoo.exceptions import UserError, ValidationError
import logging
from pypdf import PdfReader
from io import BytesIO
import pytesseract
from pdf2image import convert_from_bytes
import json
from docx import Document
from PIL import Image
import base64
import requests
import re

_logger = logging.getLogger(__name__)


class PurchaseOrder(models.Model):
    _inherit = "purchase.order"
    
    def create_document_from_attachment(self, attachment_ids):
        purchase_orders = self._create_document_from_attachment(attachment_ids)
        action_vals = {
            'name': _('Generated Purchase Orders'),
            'domain': [('id', 'in', purchase_orders.ids)],
            'res_model': 'purchase.order',
            'type': 'ir.actions.act_window',
            'context': self._context
        }
        if len(purchase_orders) == 1:
            action_vals.update({
                'views': [[False, "form"]],
                'view_mode': 'form',
                'res_id': purchase_orders[0].id,
            })
        else:
            action_vals.update({
                'views': [[False, "list"], [False, "kanban"], [False, "form"]],
                'view_mode': 'list, kanban, form',
            })
        return action_vals

    def _create_document_from_attachment(self, attachment_ids):
        attachments = self.env['ir.attachment'].browse(attachment_ids)
        if not attachments:
            raise UserError(_("No attachment was provided"))
        
        qwen_api_key = self.env['ir.config_parameter'].sudo().get_param('sttl_purchase_ocr.ocr_qwen_api_key')
        if not qwen_api_key:
            raise ValidationError(_("Please Enter valid Qwen API Key!!"))
            
        purchase_orders = self
        for attachment in attachments:
            self.action_ocr_resume(attachment, qwen_api_key)
            purchase_order = self.action_save(attachment)
            if purchase_order:
                purchase_orders += purchase_order
            
        return purchase_orders

    def action_ocr_resume(self, attachment_id, qwen_api_key):
        purchase_order = None
        for selected_record_id in attachment_id:
            binary = base64.b64decode(selected_record_id.datas)
            file_type = selected_record_id.mimetype
            text_content = ""
            try:
                if file_type == "application/pdf":
                    try:
                        pdf_reader = PdfReader(BytesIO(binary))
                        for page in pdf_reader.pages:
                            page_text = page.extract_text()
                            text_content += page_text
                        if not text_content:
                            images = convert_from_bytes(binary, dpi=300)
                            extracted_text = []
                            for image in images:
                                text = pytesseract.image_to_string(image)
                                extracted_text.append(text)
                            text_content = "\n".join(extracted_text)

                    except Exception as e:
                        _logger.error("Error reading PDF: %s", str(e))
                        text_content = ""

                elif file_type in ["image/png", "image/jpeg", "image/jpg"]:
                    try:
                        image = Image.open(BytesIO(binary))
                        text_content = pytesseract.image_to_string(image)
                    except Exception as e:
                        _logger.error("Error processing image: %s", str(e))
                        text_content = ""

                elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    try:
                        doc = Document(BytesIO(binary))
                        text_content = ''

                        for section in doc.sections:
                            header = section.header
                            for paragraph in header.paragraphs:
                                text_content += paragraph.text + '\n'

                        for paragraph in doc.paragraphs:
                            text_content += paragraph.text + '\n'
                    except Exception as e:
                        _logger.error("Error processing DOCX: %s", str(e))
                        text_content = ""

                else:
                    _logger.error("Unsupported file type: %s", file_type)
                    self.ocr_json = f"Unsupported file type: {file_type}"
                    continue
                json_response = self.get_json_from_model(text_content, qwen_api_key)

                if json_response and "choices" in json_response and len(json_response["choices"]) > 0:
                    message_content = json_response["choices"][0].get("message", {}).get("content", "")

                    if message_content:
                        match = re.search(r'```json\n(.*?)\n```', message_content, re.DOTALL)
                        if match:
                            clean_json_str = match.group(1).strip()  # Extract JSON content
                            try:
                                parsed_json = json.loads(clean_json_str)
                                selected_record_id.ocr_json = json.dumps(parsed_json, indent=4)

                            except json.JSONDecodeError as e:
                                _logger.error("Error parsing JSON: %s", str(e))
                                selected_record_id.ocr_json = "Error parsing JSON"
                        else:
                            _logger.error("No valid JSON found in the content")
                            selected_record_id.ocr_json = "No valid JSON format found"
                    else:
                        _logger.error("No message content found in the response")
                        selected_record_id.ocr_json = "No message content found"
                else:
                    _logger.error("No valid response or choices in the API response")
                    selected_record_id.ocr_json = "No valid JSON data received."

            except Exception as e:
                _logger.error("Unexpected error during OCR processing: %s", str(e))
                selected_record_id.ocr_json = "An unexpected error occurred during file processing."

            _logger.info("Stored JSON data for file: %s", selected_record_id)

        return purchase_order
    
    def get_json_from_model(self, text_content, qwen_api_key):
        api_url = "https://api.together.xyz/v1/chat/completions"
        
        headers = {
            'Authorization': 'Bearer %s' % qwen_api_key,
            'Content-Type': 'application/json',
        }
        payload = {
            "messages": [
                {
                    "role": "system",
                    "content": "Extract and format the bill details as JSON. Ensure that: \n\n"
                               "- Vendor: only take The name of the vendor mentioned in bill as vendor else dont **mostly on right upper side** (e.g., Vendor, Swing Vendor, Vendor Costing, etc). \n"
                               "- Bill Reference: Extract the **actual reference number** mentioned as 'Reference' in the bill.\n"
                               "      - If both 'Bill Reference' and 'Reference' exist, extract **'Reference'** as the main reference number.\n"
                               "- Place of supply: The location of supply mentioned in the bill if any. \n"
                               "- Bill Date: The bill issue date in '%d/%m/%Y' format. sometimes explicitely mentioned as 'Order Date' or 'Bill Date', or 'Order Deadline' etc \n"
                               "- Payment Terms: Extract **payment terms** if mentioned. \n"
                               "- Due Date: Only extract the due date if it is explicitly mentioned as 'Due Date' or 'Expected Arrival' or 'Expected' in the bill else do not extract the due date.\n"  
                               "- If 'Due Date' is missing but 'Payment Terms' exists, extract 'Payment Terms' instead.\n"
                               "- Currency: Extract the **currency symbol or code** from the bill (e.g., '£' should be converted to 'GBP'). \n"
                               "  - If the currency symbol (e.g., '£') is present, extract the correct ISO code (e.g., GBP for £, USD for $, EUR for €).\n"
                               "- Products: Extract product details with keys:\n"
                               "   - 'Product name': Name of the product.\n"
                               "   - 'Discount': Discount on the product (if available).\n"
                               "   - 'Quantity': Quantity of product.(in float)\n"
                               "   - 'Price': Price per unit.\n"
                               "   - 'Amount': Total amount before tax.\n"
                               "  - **'Tax'**: Extract the tax structure exactly as mentioned in the bill.\n"
                               "    - If the bill explicitly mentions a **parent tax** (e.g., 'GST 5%', 'VAT 20%', 'HST 13%'), extract it as a single entry without splitting.\n"
                               "    - If the bill only contains **component taxes** (e.g., 'CGST 2.5%', 'SGST 2.5%') **but no parent tax**, extract them separately.\n"
                               "    - If only tax amounts are provided (e.g., 'SGST ₹6.25', 'CGST ₹6.25'), calculate the percentage as: **(tax_amount / taxable_amount) * 100**.\n"
                               "    - Ensure tax values are always in percentage form.\n"
                               "    - Provide the tax in dictionary format with `tax_name` as key and `tax_name & tax_percentage` as value.\n"
                               "- Account Number: The account number if mentioned, else return None.\n\n"
                },
                {
                    "role": "user",
                    "content": text_content
                }
            ],
            "model": "Qwen/Qwen2.5-72B-Instruct-Turbo",
        }

        try:
            response = requests.post(api_url, json=payload, headers=headers)

            if response.status_code == 200:
                try:
                    json_response = response.json()
                    return json_response
                except ValueError as e:
                    _logger.error("Error parsing JSON: %s", str(e))
                    return {}
            else:
                _logger.error("Error in API call: %s", response.text)
                return {}

        except Exception as e:
            _logger.error("Exception during API call: %s", str(e))
            return {}


    def action_save(self, attachment_id):
        created_purchase_orders = []
        vendor = None
        new_purchase_order = None
        payment_term_id = False

        for selected_record_id in attachment_id:
            if selected_record_id.ocr_json:
                try:
                    parsed_json = json.loads(selected_record_id.ocr_json)

                    vendor_name = parsed_json.get("Vendor", "")
                    bill_reference = parsed_json.get("Bill Reference", "")
                    date_order = parsed_json.get("Bill Date", False)
                    date_planned = parsed_json.get("Due Date", False)
                    payment_term_name = parsed_json.get("Payment Terms", "")
                    # account_number = parsed_json.get("Account Number", "")
                    discount = parsed_json.get("Discount", "")
                    products = parsed_json.get("Products", [])
                    bill_currency_code = parsed_json.get("Currency", "").strip() if parsed_json.get("Currency") else None

                    # place_of_supply = parsed_json.get("Place of supply", "")

                    # company_country = self.env.company.country_id
                    # state_id = None
                    # if company_country.code == "IN" and place_of_supply:
                    #     state = self.env['res.country.state'].search([
                    #         ('name', '=', place_of_supply),
                    #         ('country_id', '=', company_country.id)
                    #     ], limit=1)
                    #     state_id = state.id if state else None


                    # Currency Selection (from bill → journal → company)
                    currency = self.env['res.currency'].search(
                        [('name', '=', bill_currency_code), ('active', '=', True)], limit=1)
                    if not currency:
                        currency = self.env['res.currency'].search(
                            [('symbol', '=', bill_currency_code), ('active', '=', True)], limit=1)
                    if not currency:
                        currency = self.env.company.currency_id

                    if not payment_term_name:
                        payment_term_name = ""
                    payment_term_name = payment_term_name.strip()

                    if date_order:
                        date_order = datetime.strptime(date_order, '%d/%m/%Y')

                    if date_planned and date_planned != None:
                        date_planned = datetime.strptime(date_planned, '%d/%m/%Y')
                    else:
                        date_planned = None

                    payment_term_id = False
                    if payment_term_name and date_planned == None:
                        payment_term = self.env['account.payment.term'].search([
                            ('name', '=', payment_term_name)
                        ], limit=1)

                        if payment_term:
                            payment_term_id = payment_term.id

                    if vendor_name:
                        vendor_name = vendor_name.strip()
                        vendor = self.env['res.partner'].search([('name', '=', vendor_name)], limit=1)

                        if not vendor:
                            vendor = self.env['res.partner'].create({'name': vendor_name})

                    if discount:
                        discount = str(discount).replace('%', '').strip()
                        discount = float(discount) if discount else 0.0
                    else:
                        discount = 0.0

                    purchase_order_lines = []
                    for product_data in products:
                        product_name = None
                        if isinstance(product_data, dict):

                            product_name = product_data.get('Product name', '')
                            quantity = product_data.get('Quantity', 0)
                            price = product_data.get('Price', 0.0)
                            tax_data = product_data.get('Tax', {})
                            discount = str(product_data.get('Discount', '')).replace('%', '') if product_data.get('Discount') else 0.0
                            # amount = product_data.get('Amount', 0.0)

                            product = self.env['product.product'].search([('name', '=', product_name)], limit=1)
                            if not product:
                                product_supplier = self.env['product.supplierinfo'].search([('product_name', '=', product_name)], limit=1)
                                if product_supplier:
                                    if product_supplier.product_id:
                                        product = product_supplier.product_id
                                    elif product_supplier.product_tmpl_id:
                                        product = self.env['product.product'].search([('product_tmpl_id', '=', product_supplier.product_tmpl_id.id)])
                                if not product:
                                    product = self.env['product.product'].create({'name': product_name})

                            taxes = []
                            tax_percentage = 0.0

                            # ✅ Ensure tax data is a dictionary
                            if isinstance(tax_data, (int, float)):
                                tax_percentage = float(tax_data)
                                tax_data = {f"GST {int(tax_percentage)}%": tax_percentage}  # Keep % format
                            elif not isinstance(tax_data, dict):
                                tax_data = {}


                            for tax_name, tax_value in tax_data.items():
                                tax_name = str(tax_name).strip()

                                # ✅ Extract numeric tax percentage if formatted like "GST 12%"
                                if isinstance(tax_value, str):
                                    match = re.search(r"(\d+)%", tax_value)
                                    if match:
                                        tax_percentage = float(match.group(1))
                                # ✅ Normalize Tax Name for better matching
                                normalized_tax_name = tax_name.upper()

                                # ✅ Possible Tax Name Variations
                                possible_tax_names = [
                                    normalized_tax_name,  # "GST 12%"
                                    f"IGST {int(tax_percentage)}%",  # "IGST 12%"
                                    f"{int(tax_percentage)}% IGST P",  # "12% IGST P"
                                    f"{int(tax_percentage)}% IGST",
                                    f"GST {int(tax_percentage)}%",  # "GST 12%"
                                    f"{int(tax_percentage)}% GST",  # "12% GST"
                                    f"Tax {int(tax_percentage)}%",  # "Tax 12%"
                                    f"GST {int(tax_percentage)}",  # "GST 12"
                                    f"{int(tax_percentage)}% GST",
                                    f"{int(tax_percentage)}%",  # "12%"
                                    f"{int(tax_percentage)}% GST P",  # "12% GST P"
                                    f"{int(tax_percentage)}% GST S",  # "12% GST S"
                                ]
                                taxes_found = self.env['account.tax'].search([
                                    ('company_id', '=', self.env.company.id),
                                    ('amount', '=', tax_percentage),
                                    ('amount_type', 'in', ['group', 'percent']),
                                    ('type_tax_use', '=', 'purchase'),
                                    '|',
                                    ('name', 'in', [f"{int(tax_percentage)}% {tax_name} P"]),
                                    ('name', 'in', [tax_name]),
                                ], order="id asc", limit=1)

                                if not taxes_found:
                                    tax_group = self.env['account.tax'].search([
                                        ('company_id', '=', self.env.company.id),
                                        ('amount', '=', tax_percentage),
                                        ('amount_type', 'in', ['group', 'percent']),
                                        ('type_tax_use', '=', 'purchase'),
                                        '|',
                                        ('name', 'in', [f"{int(tax_percentage)}% {tax_name}"]),
                                        ('name', 'in', [tax_name]),
                                    ], order="id asc", limit=1)
                                    if tax_group:
                                        taxes_found = tax_group

                                if not taxes_found:
                                    tax_group = self.env['account.tax'].search([
                                        ('company_id', '=', self.env.company.id),
                                        ('amount', '=', tax_percentage),
                                        ('amount_type', 'in', ['group', 'percent']),
                                        ('type_tax_use', '=', 'purchase'),
                                        '|',
                                        ('name', 'ilike', f"{int(tax_percentage)}% {tax_name}"),
                                        ('name', 'ilike', tax_name),
                                        ('name', 'in', possible_tax_names),
                                    ], order="id asc", limit=1)
                                    if tax_group:
                                        taxes_found = tax_group

                                if not taxes_found:
                                    tax_group = self.env['account.tax'].search([
                                        ('company_id', '=', self.env.company.id),
                                        ('amount', '=', tax_percentage),
                                        ('amount_type', '=', 'group'),
                                        ('name', 'ilike', f"%{int(tax_percentage)}%"),
                                    ], limit=1)

                                    if tax_group:
                                        taxes_found = tax_group
                                        
                                if not taxes_found:
                                    tax_group = self.env['account.tax'].search([
                                        ('company_id', '=', self.env.company.id),
                                        ('name', '=', tax_value),
                                        ('amount_type', 'in', ['group', 'percent']),
                                    ], limit=1)

                                    if tax_group:
                                        taxes_found = tax_group

                                if taxes_found:
                                    taxes.append(taxes_found.id)
                                else:
                                    _logger.warning(
                                        f"Warning: No existing tax found for {tax_name} {tax_percentage}%")

                            taxes = list(set(taxes))

                            if not taxes:
                                taxes = []

                            purchase_order_lines.append((0, 0, {
                                'product_id': product.id,
                                'product_qty': quantity,
                                'discount': discount,
                                'price_unit': price,
                                'taxes_id': [(6, 0, taxes)] if taxes else [],
                            }))

                    # Find or create a bank account for the vendor
                    # if account_number:
                    #     partner_bank = self.env['res.partner.bank'].search([('partner_id', '=', vendor.id)], limit=1)
                    #     if not partner_bank:
                    #         partner_bank = self.env['res.partner.bank'].create({
                    #             'acc_number': account_number,
                    #             'partner_id': vendor.id if vendor else False,
                    #             'bank_id': self.env['res.bank'].search([], limit=1).id,
                    #         })
                    #     partner_bank_id = partner_bank.id
                    # else:
                    #     partner_bank_id = False

                    new_purchase_order = self.env['purchase.order'].create({
                        # 'move_type': 'in_invoice',
                        'partner_id': vendor.id if vendor else False,
                        'date_order':  date_order if date_order else fields.Date.today(),
                        # 'partner_bank_id': partner_bank_id,
                        'partner_ref': bill_reference,
                        # 'is_ocr': True,
                        'date_planned':  date_planned if date_planned else False,
                        'payment_term_id': payment_term_id,
                        'currency_id': currency.id,
                        # 'journal_id': journal.id,
                        'order_line': purchase_order_lines,
                        # 'date' : fields.Date.today(),
                    })

                    # if company_country.code == "IN":
                    #     new_purchase_order['l10n_in_state_id'] = state_id

                    if new_purchase_order and selected_record_id:
                        selected_record_id.write({
                            'res_model': 'purchase.order',
                            'res_id': new_purchase_order.id
                        })

                    created_purchase_orders.append(new_purchase_order.id)

                except json.JSONDecodeError as e:
                    _logger.error("Error parsing JSON data for file %s: %s", selected_record_id.name, str(e))
            else:
                _logger.warning("No JSON data found for file %s", selected_record_id.name)
        return new_purchase_order if new_purchase_order else self.env['purchase.order']

    